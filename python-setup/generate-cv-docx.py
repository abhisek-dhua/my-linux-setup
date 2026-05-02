from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Create a new one-page styled resume with full Project section and compacted others
doc = Document()
for section in doc.sections:
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

# Header
header = doc.add_paragraph()
header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = header.add_run("Abhisek Dhua")
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0, 51, 102)

sub_header = doc.add_paragraph()
sub_header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = sub_header.add_run("Angular Developer | +91 6294373636 | abhisek.dhua@gmail.com")
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(80, 80, 80)

def add_title(doc, title):
    p = doc.add_paragraph()
    run = p.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0, 102, 204)
    p.space_before = Pt(6)
    p.space_after = Pt(2)

def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(item)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(30, 30, 30)

def add_paragraph(doc, texts):
    for text in texts:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(30, 30, 30)

# Professional Summary (shortened slightly)
add_title(doc, "Professional Summary")
add_paragraph(doc, [
    "Angular Developer with 3+ years' experience in building scalable, high-performance web apps. Skilled in translating complex requirements into efficient, user-centric solutions. Proven success across e-commerce, CMS, and enterprise systems. Strong team player committed to code quality, mentoring, and agile delivery."
])

# Skills (single-line, condensed)
add_title(doc, "Skills")
add_paragraph(doc, [
    "Angular (v8–v15+), RxJS, HTML5, CSS3, JavaScript, TypeScript, SCSS, Tailwind, Angular Material, RESTful APIs, GraphQL, JSON, Git, GitHub, Bitbucket, Postman, VS Code, Agile, Scrum, Problem-Solving, Mentoring."
])

# Professional Experience (compressed)
add_title(doc, "Professional Experience")
add_paragraph(doc, ["Senior Web Developer | Mass Software Solutions Pvt Ltd | 2022 - Present"])
add_bullets(doc, [
    "Developed scalable Angular apps with best practices and performance techniques (lazy loading, refactoring).",
    "Integrated complex APIs, improved UX, and maintained documentation.",
    "Mentored junior devs and collaborated with cross-functional teams.",
    "Wrote tests (unit/E2E) and contributed to Agile ceremonies."
])

# Projects (unchanged as requested)
add_title(doc, "Projects")
add_bullets(doc, [
    "School Food Delivery System (Ireland): Developed key front-end features for an online platform for ordering and delivering school meals, and a strong admin panel to manage products, menus, and the entire back-end.",
    "E-commerce Platform: Contributed to building a responsive and feature-rich e-commerce website with a focus on user experience and seamless navigation.",
    "Insurance Project: Worked on a web application for managing insurance policies, claims, and customer data with a secure and intuitive interface.",
    "Legal Advice Platform: Developed UI components for a platform connecting users with legal professionals, ensuring confidentiality and ease of use.",
    "Advanced Charging Station Project: Built the front-end interface for a system managing and monitoring electric vehicle charging stations.",
    "Watchmaking Company Project: Designed and implemented an inventory management system for a watchmaking company to track parts, ordering parts and sales.",
    "CMS System: Developed a custom Content Management System to empower non-technical users to manage website content efficiently."
])

# Education (compact format)
add_title(doc, "Education")
add_paragraph(doc, [
    "B.Sc. in Computer Science, University of Burdwan | 2017 | 53%",
    "12th Grade, Bhedua Salboni Naba Siksha Mandir H.S. | 2014 | 64.4%"
])

# Save final one-page version
final_path = "Abhisek_Dhua_Resume_One_Page_Final.docx"
doc.save(final_path)


print("📄 DOCX file generated successfully!")
print("💡 To convert to PDF, use LibreOffice: libreoffice --headless --convert-to pdf Abhisek_Dhua_Resume_One_Page_Final.docx")